"""多格式文档加载器 - 支持PDF/Word/Markdown/CSV/JSON"""

import json
from typing import List, Union
from pathlib import Path
from langchain_core.documents import Document
import re


class SmartDocumentLoader:
    """智能多格式文档加载器"""

    def __init__(self, file_path: Union[str, Path]):
        self.file_path = Path(file_path)
        self.ext = self.file_path.suffix.lower()

    def load(self) -> List[Document]:
        """根据文件类型自动选择加载器"""
        loaders = {
            ".pdf": self._load_pdf,
            ".docx": self._load_word,
            ".doc": self._load_word,
            ".md": self._load_markdown,
            ".txt": self._load_text,
            ".csv": self._load_csv,
            ".html": self._load_html,
            ".json": self._load_json,
        }

        loader = loaders.get(self.ext, self._load_text)
        return loader()

    def _load_pdf(self) -> List[Document]:
        """PDF加载（保留布局信息）"""
        try:
            from langchain_community.document_loaders import PyMuPDFLoader

            loader = PyMuPDFLoader(str(self.file_path))
            docs = loader.load()

            for i, doc in enumerate(docs):
                doc.metadata.update(
                    {
                        "page_number": doc.metadata.get("page", i),
                        "source_type": "pdf",
                        "file_name": self.file_path.name,
                        "total_pages": len(docs),
                    }
                )
            return docs
        except ImportError:
            # 降级方案
            from langchain_community.document_loaders import PyPDFLoader

            return PyPDFLoader(str(self.file_path)).load()

    def _load_word(self) -> List[Document]:
        """Word文档（提取表格和标题）"""
        from langchain_community.document_loaders import Docx2txtLoader

        loader = Docx2txtLoader(str(self.file_path))
        docs = loader.load()

        # 尝试提取Word表格
        try:
            from docx import Document as DocxDocument

            docx = DocxDocument(self.file_path)
            tables_text = []

            for table in docx.tables:
                table_data = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_data.append("| " + " | ".join(row_text) + " |")
                if table_data:
                    table_text = "\n".join(table_data)
                    records_text = self._table_markdown_to_records(table_text)
                    if records_text:
                        table_text = f"{table_text}\n\n[表格记录]\n{records_text}"
                    tables_text.append(table_text)

            if docs and tables_text:
                docs[0].page_content += "\n\n[文档中的表格]:\n" + "\n\n".join(
                    tables_text
                )
                docs[0].metadata["has_tables"] = True
                docs[0].metadata["table_count"] = len(tables_text)
        except Exception:
            pass

        for doc in docs:
            doc.metadata.update(
                {"source_type": "word", "file_name": self.file_path.name}
            )
        return docs

    def _load_markdown(self) -> List[Document]:
        """Markdown（保留Frontmatter和目录结构）"""
        try:
            import yaml
        except ImportError:
            yaml = None

        with open(self.file_path, "r", encoding="utf-8") as f:
            content = f.read()

        metadata = {"source_type": "markdown", "file_name": self.file_path.name}

        # 解析YAML Frontmatter
        if yaml:
            frontmatter_match = re.match(r"^---\n(.*?)\n---\n", content, re.DOTALL)
            if frontmatter_match:
                try:
                    front_data = yaml.safe_load(frontmatter_match.group(1))
                    if isinstance(front_data, dict):
                        metadata.update(front_data)
                    content = content[frontmatter_match.end() :]
                except Exception:
                    pass

        # 提取目录结构
        headers = re.findall(r"^(#{1,3})\s+(.+)$", content, re.MULTILINE)
        metadata["outline"] = json.dumps(
            [{"level": len(h[0]), "title": h[1]} for h in headers[:10]]
        )

        return [Document(page_content=content, metadata=metadata)]

    def _load_csv(self) -> List[Document]:
        """CSV作为表格文档处理"""
        import pandas as pd

        df = pd.read_csv(self.file_path)

        # 转换为Markdown表格
        markdown_table = df.to_markdown(index=False)
        summary = f"表格包含 {len(df)} 行, {len(df.columns)} 列。列名: {', '.join(df.columns)}"

        # 数据类型摘要（转为字符串而非列表）
        dtypes = df.dtypes.to_dict()
        type_summary = json.dumps(
            {k: str(v) for k, v in dtypes.items()}
        )  # 转为JSON字符串

        records_text = self._dataframe_to_records(df)

        return [
            Document(
                page_content=f"{summary}\n\n{markdown_table}\n\n[表格记录]\n{records_text}",
                metadata={
                    "source_type": "csv",
                    "file_name": self.file_path.name,
                    "row_count": len(df),
                    "columns": ", ".join(df.columns),  # 改为逗号分隔的字符串
                    "dtypes": type_summary,  # JSON字符串
                },
            )
        ]

    def _table_markdown_to_records(
        self, markdown_table: str, limit_rows: int = 200
    ) -> str:
        """将Markdown表格转为通用记录文本（字段=值），适配任意字段名。"""
        lines = [
            line.strip()
            for line in markdown_table.splitlines()
            if line.strip().startswith("|")
        ]
        if len(lines) < 2:
            return ""

        rows = []
        for line in lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)

        headers = rows[0]
        data_rows = []
        for r in rows[1:]:
            # 跳过分隔行
            if all(re.fullmatch(r"[:\- ]*", c or "") for c in r):
                continue
            data_rows.append(r)

        if not headers or not data_rows:
            return ""

        normalized_headers = [h if h else f"字段{i + 1}" for i, h in enumerate(headers)]
        record_lines = []
        for idx, row in enumerate(data_rows[:limit_rows], 1):
            pairs = []
            for i, header in enumerate(normalized_headers):
                value = row[i].strip() if i < len(row) else ""
                pairs.append(f"{header}={value}")
            record_lines.append(f"记录{idx}: " + " ; ".join(pairs))
        return "\n".join(record_lines)

    def _dataframe_to_records(self, df, limit_rows: int = 500) -> str:
        """DataFrame通用记录化，避免只靠Markdown表格检索。"""
        if df.empty:
            return ""

        safe_df = df.fillna("")
        cols = [
            str(c).strip() if str(c).strip() else f"字段{i + 1}"
            for i, c in enumerate(safe_df.columns)
        ]
        record_lines = []
        for idx, row in enumerate(safe_df.itertuples(index=False), 1):
            values = [str(v).strip() for v in row]
            pairs = [f"{k}={v}" for k, v in zip(cols, values)]
            record_lines.append(f"记录{idx}: " + " ; ".join(pairs))
            if idx >= limit_rows:
                break
        return "\n".join(record_lines)

    def _load_html(self) -> List[Document]:
        """HTML（清理标签，提取正文）"""
        from langchain_community.document_loaders import BSHTMLLoader

        loader = BSHTMLLoader(str(self.file_path), open_encoding="utf-8")
        docs = loader.load()

        for doc in docs:
            doc.metadata.update(
                {"source_type": "html", "file_name": self.file_path.name}
            )
        return docs

    def _load_text(self) -> List[Document]:
        """纯文本"""
        with open(self.file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return [
            Document(
                page_content=content,
                metadata={
                    "source_type": "text",
                    "file_name": self.file_path.name,
                    "char_count": len(content),
                },
            )
        ]

    def _load_json(self) -> List[Document]:
        """JSON（支持JSONL和结构化JSON）"""
        import json

        with open(self.file_path, "r", encoding="utf-8") as f:
            content = f.read()
            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                # 尝试JSONL格式
                f.seek(0)
                data = [json.loads(line) for line in f if line.strip()]

        if isinstance(data, list):
            # 列表格式：每个元素一个文档
            return [
                Document(
                    page_content=json.dumps(item, ensure_ascii=False, indent=2),
                    metadata={
                        "source_type": "json",
                        "index": i,
                        "file_name": self.file_path.name,
                    },
                )
                for i, item in enumerate(data[:100])  # 限制数量
            ]
        else:
            # 单个对象
            return [
                Document(
                    page_content=json.dumps(data, ensure_ascii=False, indent=2),
                    metadata={"source_type": "json", "file_name": self.file_path.name},
                )
            ]


class BatchDocumentLoader:
    """批量加载文件夹"""

    def __init__(self, directory: str, glob_pattern: str = "**/*"):
        self.directory = Path(directory)
        self.pattern = glob_pattern

    def load(self) -> List[Document]:
        """批量加载目录下所有支持的文件"""
        documents = []
        supported = [".pdf", ".docx", ".doc", ".md", ".txt", ".csv", ".html", ".json"]

        for file_path in self.directory.glob(self.pattern):
            if file_path.is_file() and file_path.suffix.lower() in supported:
                try:
                    loader = SmartDocumentLoader(file_path)
                    docs = loader.load()
                    documents.extend(docs)
                    print(f"[Loader] 加载: {file_path}")
                except Exception as e:
                    print(f"[Loader] 失败: {file_path} - {e}")

        return documents
